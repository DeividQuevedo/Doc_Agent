import os
import yaml
from langchain_community.vectorstores import FAISS
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.schema import Document
from docx import Document as DocxDocument
from PyPDF2 import PdfReader

# Carregar configuração da API
with open('config.yaml', 'r') as config_file:
    config = yaml.safe_load(config_file)
os.environ['OPENAI_API_KEY'] = config['OPENAI_API_KEY']

def extract_text_from_txt(file_path):
    """Extrai texto de arquivos .txt."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def extract_text_from_docx(file_path):
    """Extrai texto de arquivos .docx."""
    doc = DocxDocument(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

def extract_text_from_pdf(file_path):
    """Extrai texto de arquivos .pdf."""
    pdf_reader = PdfReader(file_path)
    return "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])

def load_documents(folder_path):
    """Carrega documentos .txt, .docx e .pdf, gera embeddings e cria o vetor FAISS."""
    documents = []
    for file_name in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file_name)

        # Processar arquivos .txt
        if file_name.endswith(".txt"):
            try:
                content = extract_text_from_txt(file_path)
                if content.strip():
                    documents.append(Document(page_content=content))
                else:
                    print(f"Aviso: O arquivo '{file_name}' está vazio.")
            except Exception as e:
                print(f"Erro ao processar o arquivo .txt '{file_name}': {e}")

        # Processar arquivos .docx
        elif file_name.endswith(".docx"):
            try:
                content = extract_text_from_docx(file_path)
                if content.strip():
                    documents.append(Document(page_content=content))
                else:
                    print(f"Aviso: O arquivo '{file_name}' está vazio.")
            except Exception as e:
                print(f"Erro ao processar o arquivo .docx '{file_name}': {e}")

        # Processar arquivos .pdf
        elif file_name.endswith(".pdf"):
            try:
                content = extract_text_from_pdf(file_path)
                if content.strip():
                    documents.append(Document(page_content=content))
                else:
                    print(f"Aviso: O arquivo '{file_name}' está vazio.")
            except Exception as e:
                print(f"Erro ao processar o arquivo .pdf '{file_name}': {e}")

        else:
            print(f"Aviso: O arquivo '{file_name}' não é suportado.")

    if not documents:
        raise ValueError("Nenhum documento válido (.txt, .docx, .pdf) encontrado no diretório especificado.")

    print(f"{len(documents)} documentos carregados.")

    # Criar embeddings e índice FAISS
    embeddings = OpenAIEmbeddings()
    try:
        vector_store = FAISS.from_documents(documents, embeddings)
    except Exception as e:
        raise RuntimeError("Erro ao criar o índice FAISS.") from e

    return vector_store
